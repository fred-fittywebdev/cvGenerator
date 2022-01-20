from turtle import width
from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(txt):
    pyttsx3.speak(txt)

document = Document()

document.add_picture("avatar.jpg", width=Inches(2.0))

speak("Hi and welcome, simply generate your resume by answering the questions, it's that easy!")

first_name = input("Entrez votre prénom: ")
speak('Hello ' + first_name + ' how are you today? ')
last_name = input("Entrez votre non de famille: ")
speak('Enter your phone number please ! ')
phone_number = input("Entrez votre numéro de téléphone: ")
speak('Enter your adress now please ! ')
adress = input("Entrez votre adresse: ")
cp_ville = input("Ville et code postal: ")
speak('One last things, enter your email adress ! ')
email = input("Entrez votre adresse mail: ")


document.add_paragraph(
    last_name + " | " + first_name + "\n" + 
    adress + "\n" + 
    cp_ville + "\n" +
    phone_number + "\n" + 
    email
)

# About section
document.add_heading("Qui suis-je? ")
about_me = input("Parlez de vous, c'est la phrase d'accroche ;-). ")
document.add_paragraph(about_me)

# Work experience
document.add_heading("Expériences professionnelles: ")
p = document.add_paragraph()

entreprise = input("Nom de l'entreprise: ")
from_date = input("de: ")
to_date = input("jusqu'au: ")

p.add_run(entreprise + ' ').bold = True  # ad_run permet d'ajouter du contenu a un paragraphe
p.add_run("De " + from_date + ' à ' + to_date + '.\n').italic = True

experience_description = input(f"Décrivez votre esperience chez {entreprise}. ")
p.add_run(experience_description)

# Add more expreiences while loop
while True:
    experiences_supplementaire = input("Voulez vous ajouter d'autres expériences professionnelles? (oui/non)").lower()
    if experiences_supplementaire == "oui":
        p = document.add_paragraph()

        entreprise = input("Nom de l'entreprise: ")
        from_date = input("De: ")
        to_date = input("jusqu'au: ")

        p.add_run(entreprise + ' ').bold = True  # ad_run permet d'ajouter du contenu a un paragraphe
        p.add_run("De " + from_date + ' à ' + to_date + '.\n').italic = True

        experience_description = input(f"Décrivez votre esperience chez {entreprise}. ")
        p.add_run(experience_description)
    else:
        break

# Skills
document.add_heading("Compétences: ")
skills = input("Entrez une compétence ")
p = document.add_paragraph(skills)
p.style = "List Bullet"

# Add more compétences while loop
while True:
    competence_supplementaire = input("Voulez vous ajouter d'autres compétences? (oui/non) ")
    if competence_supplementaire.lower() == "oui":
        skills = input("Entrez une compétence ")
        p.style = "List Bullet"
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV généré grâce a un script python rendre la création de CV un peu plus fun"

document.save("Cv.docx")